# -*- coding: utf-8 -*-
"""
将 Markdown 简历编译为专业排版的 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re


class ResumeCompiler:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """设置文档样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

    def add_name_header(self, name, contact_info):
        """添加姓名和联系信息"""
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        name_para.paragraph_format.space_after = Pt(6)

        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, info in enumerate(contact_info):
            run = contact_para.add_run(info)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if i < len(contact_info) - 1:
                contact_para.add_run(" | ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        contact_para.paragraph_format.space_after = Pt(12)

    def add_section_title(self, title):
        """添加章节标题"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="1a1a2e"/></w:pBdr>')
        para._p.get_or_add_pPr().append(pBdr)

    def add_education(self, school, major, degree, start_date, end_date):
        """添加教育背景"""
        self.add_section_title("EDUCATION")

        para = self.doc.add_paragraph()
        run = para.add_run(f"{school}")
        run.font.bold = True
        run.font.size = Pt(11)
        para.add_run(f" — {major} | {degree}").font.size = Pt(10)

        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"{start_date} - {end_date}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_para.paragraph_format.space_after = Pt(4)

    def add_skills(self, skills_dict):
        """添加技能部分"""
        self.add_section_title("TECHNICAL SKILLS")

        for category, skills in skills_dict.items():
            para = self.doc.add_paragraph()
            cat_run = para.add_run(f"{category}: ")
            cat_run.font.bold = True
            cat_run.font.size = Pt(10)

            skills_run = para.add_run(", ".join(skills))
            skills_run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)

    def add_experience(self, company, position, start_date, end_date, highlights):
        """添加项目经历"""
        header_para = self.doc.add_paragraph()

        company_run = header_para.add_run(company)
        company_run.font.bold = True
        company_run.font.size = Pt(11)

        header_para.add_run(" | ").font.size = Pt(10)

        position_run = header_para.add_run(position)
        position_run.font.size = Pt(10)

        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"{start_date} - {end_date}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(4)

        if "description" in highlights:
            desc_para = self.doc.add_paragraph()
            desc_run = desc_para.add_run(highlights["description"])
            desc_run.font.italic = True
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            desc_para.paragraph_format.space_after = Pt(4)

        if "tech_stack" in highlights:
            tech_para = self.doc.add_paragraph()
            tech_run = tech_para.add_run("Tech Stack: ")
            tech_run.font.bold = True
            tech_run.font.size = Pt(9)

            tech_detail_run = tech_para.add_run(highlights["tech_stack"])
            tech_detail_run.font.size = Pt(9)
            tech_detail_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            tech_para.paragraph_format.space_after = Pt(4)

        if "bullets" in highlights:
            for bullet in highlights["bullets"]:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.size = Pt(10)
                bullet_para.paragraph_format.space_after = Pt(2)

        if "metrics" in highlights:
            metrics_para = self.doc.add_paragraph()
            metrics_title_run = metrics_para.add_run("量化压测成果:")
            metrics_title_run.font.bold = True
            metrics_title_run.font.size = Pt(10)
            metrics_para.paragraph_format.space_after = Pt(2)

            for metric in highlights["metrics"]:
                metric_para = self.doc.add_paragraph(style='List Bullet')
                metric_run = metric_para.add_run(metric)
                metric_run.font.size = Pt(10)
                metric_para.paragraph_format.space_after = Pt(2)

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)
        print(f"简历已保存为: {filename}")


def create_optimized_resume():
    """创建优化后的简历"""
    compiler = ResumeCompiler()

    # 联系信息
    compiler.add_name_header(
        name="王辰宇",
        contact_info=[
            "电话: 13797390580",
            "邮箱: 2040188027@qq.com",
            "GitHub: [项目已开源，面试时提供远程查阅]"
        ]
    )

    # 教育背景
    compiler.add_education(
        school="武昌首义学院",
        major="计算机科学与技术（华为云班）",
        degree="本科在读",
        start_date="2023.09",
        end_date="2027.06"
    )

    # 专业技能
    skills = {
        "AI 后端与并发": [
            "Python asyncio 非阻塞编程",
            "协程并发与线程池隔离调优",
            "HTTP SSE 流式传输",
            "TCP 背压机制"
        ],
        "RAG 与检索优化": [
            "双阶段混合召回 (Dense/Sparse)",
            "RRF 倒数排名融合",
            "Cross-Encoder 精排",
            "向量语义缓存 (VSS)"
        ],
        "Agent 与智能体": [
            "状态机 (FSM) 架构",
            "ReAct 确定性运行时",
            "Tool Calling 机制",
            "任务规划与迭代控制"
        ],
        "文档解析与 ETL": [
            "多模态文档解析",
            "长上下文切分策略",
            "表格防断裂机制",
            "VLM 图像语义提取"
        ],
        "云原生与中间件": [
            "FastAPI",
            "Redis / Qdrant / Neo4j",
            "Celery 分布式调度",
            "Docker / MinIO"
        ]
    }
    compiler.add_skills(skills)

    # 项目经历 1: KGateway
    compiler.add_experience(
        company="KGateway - 企业级混合多模态 Agent 知识库网关",
        position="AI 后端研发",
        start_date="2026.04",
        end_date="2026.05",
        highlights={
            "description": "面向企业大模型落地场景的 AI 基础设施网关，基于 FastAPI 构建流式异步管道，解决长尾延迟、Token 成本高及多租户隔离问题。",
            "tech_stack": "FastAPI, Qdrant, Redis, Neo4j, BGE-Reranker, LangFuse, Locust",
            "bullets": [
                "设计 HTTP SSE 双向解耦流式网关，实现「状态流 + 文本流」架构；基于 asyncio.wait(FIRST_COMPLETED) 实现双路竞速守护机制，在模型思考期维持 200ms 高频客户端心跳检测。检测到客户端离线后，显式调用 pending_task.cancel() 并通过 CancelledError 异常传导链显式回收 pending 协程，强杀上游 HTTP/2 RST_STREAM 帧，从根本上杜绝孤儿 Task 驻留事件循环引发的内存泄漏，毫秒级终止上游计费，减少无效 Token 消耗。",
                "基于 Qdrant HNSW 索引原生 Filter 实现 AND 逻辑预过滤位图剪枝，在索引遍历阶段完成租户数据硬隔离，规避后过滤导致的召回率衰减问题。",
                "构建 Dense/Sparse 双阶段混排管线：采用 2-gram 倒排索引 + RRF 融合算法（k=60）动态排序；通过 asyncio.to_thread 安全卸载 BGE-Reranker 精排推理，避免阻塞主事件循环。",
                "摒弃黑盒框架，基于 FSM 状态机自研确定性 Agent 运行时（Planner → Tool Executor → Fallback），注入最大 4 次迭代沙箱，显著降低无限循环风险。",
                "自研 Closed/Open/Half-Open 三态自修复熔断器，基于 60s 滑动窗口（10次错误率阈值）动态检测下游 API 状态；在 Half-Open 状态下通过单并发请求进行动态探测自修复，防止高并发下游 API 速率限制（429）雪崩蔓延。",
                "实现 RediSearch 向量语义缓存，对语义相似度 >0.96 的请求实现 5ms 内无损拦截，显著降低 Token 财务成本。"
            ],
            "metrics": [
                "经 Locust 分布式压力测试验证，在支持 200 并发 SSE 长连接的极限压测基准下，热路径命中向量语义缓存（VSS）实现 5ms 内无损拦截，冷路径 P99 首包延迟（TTFT）稳定压制在 45ms 内。",
                "在 1000 QPS 压测基准下，基于 1000 条真实测试集调参确立 0.96 向量相似度阈值，语义缓存最高命中率达 68.4%，使单日 Token 财务成本降低 42.6%。"
            ]
        }
    )

    # 项目经历 2: OmniParse
    compiler.add_experience(
        company="OmniParse - 企业级异步多模态文档解析与 ETL 引擎",
        position="AI 数据架构",
        start_date="2026.03",
        end_date="2026.04",
        highlights={
            "description": "面向大模型异构知识库场景的分布式数据清洗与入库流水线，支持复杂 PDF、跨页表格等多源数据，提升脏数据结构化提取的准确率。",
            "tech_stack": "Celery, Redis, MinIO, Unstructured, Qdrant, Docker Compose",
            "bullets": [
                "设计 FastAPI + Celery 分布式异步解耦架构，将耗时文件解析从主业务分离。将轻量级 VLM 推理以进程级常驻方式部署于独立算力队列 Worker，通过配置 CELERYD_MAX_TASKS_PER_CHILD=50 定期强行回收子进程，规避显存碎片累积引发的 CUDA OOM 风险。",
                "自研格式感知切分器 EnterpriseChunker，支持大文件分片传入推理队列并限制单批次图片数量，将空间复杂度稳稳压制在 O(1)。在切分阶段将跨页表格整体还原为完整的 Markdown/HTML 结构作为 Parent Doc，利用轻量级 LLM 生成摘要作为 Child Chunk 入库并保留指针（父子文档策略），显著降低因语义截断引发的模型幻觉率。",
                "在结构化解析阶段透传 PDF 内嵌图片至轻量级 VLM 生成 Image Caption，完成富文本原位合并，保留长上下文图文关联性。",
                "设计全流式 Generator Pipeline，将 Qdrant 入库控制在 100 points/batch，高并发场景空间复杂度压制至 O(1)；基于 Docker Compose 实现多级 Worker 一键部署。",
                "高并发场景下任务提交 P99 延迟 < 20ms。"
            ]
        }
    )

    # 保存文件（先保存为临时文件，再重命名）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    temp_path = os.path.join(script_dir, "Resume_new.docx")
    final_path = os.path.join(script_dir, "Resume.docx")

    compiler.save(temp_path)

    # 尝试重命名替换旧文件
    import time
    time.sleep(0.5)  # 等待文件释放

    try:
        if os.path.exists(final_path):
            os.remove(final_path)
        os.rename(temp_path, final_path)
        print(f"已成功覆盖: {final_path}")
    except Exception as e:
        print(f"无法覆盖原文件，新文件已保存为: {temp_path}")
        print(f"请手动关闭旧文件后重命名。错误信息: {e}")


if __name__ == "__main__":
    create_optimized_resume()
    print("简历编译完成！")
