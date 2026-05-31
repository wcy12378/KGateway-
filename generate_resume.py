# -*- coding: utf-8 -*-
"""
简历生成器 - 生成专业排版的 Word 简历
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


class ResumeGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 设置页边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

    def add_name_header(self, name, phone, email, github=None, linkedin=None):
        """添加姓名和联系信息头部"""
        # 姓名
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        name_para.paragraph_format.space_after = Pt(6)

        # 联系信息
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(f"{phone}  |  {email}")
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if github:
            contact_para.add_run(f"  |  GitHub: {github}").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        contact_para.paragraph_format.space_after = Pt(12)

    def add_section_title(self, title):
        """添加章节标题"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # 添加下划线效果
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

        # 添加底部边框
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="1a1a2e"/></w:pBdr>')
        para._p.get_or_add_pPr().append(pBdr)

    def add_education(self, school, major, degree, start_date, end_date, details=None):
        """添加教育背景"""
        self.add_section_title("EDUCATION")

        # 学校和专业
        para = self.doc.add_paragraph()
        run = para.add_run(f"{school}")
        run.font.bold = True
        run.font.size = Pt(11)

        para.add_run(f"  |  {major}  |  {degree}").font.size = Pt(10)

        # 日期
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"{start_date} - {end_date}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_para.paragraph_format.space_after = Pt(4)

    def add_skills(self, skills_dict):
        """添加技能部分"""
        self.add_section_title("TECHNICAL SKILLS")

        for category, skills in skills_dict.items():
            para = self.doc.add_paragraph()
            cat_run = para.add_run(f"{category}: ")
            cat_run.font.bold = True
            cat_run.font.size = Pt(10)

            skills_run = para.add_run(", ".join(skills))
            skills_run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)

    def add_experience(self, company, position, start_date, end_date, highlights):
        """添加工作/项目经历"""
        # 公司和职位
        header_para = self.doc.add_paragraph()

        company_run = header_para.add_run(company)
        company_run.font.bold = True
        company_run.font.size = Pt(11)

        header_para.add_run("  |  ").font.size = Pt(10)

        position_run = header_para.add_run(position)
        position_run.font.size = Pt(10)

        # 日期（右对齐通过制表符）
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"{start_date} - {end_date}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(4)

        # 项目简介
        if "description" in highlights:
            desc_para = self.doc.add_paragraph()
            desc_run = desc_para.add_run(highlights["description"])
            desc_run.font.italic = True
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            desc_para.paragraph_format.space_after = Pt(4)

        # 技术栈
        if "tech_stack" in highlights:
            tech_para = self.doc.add_paragraph()
            tech_run = tech_para.add_run("Tech Stack: ")
            tech_run.font.bold = True
            tech_run.font.size = Pt(9)

            tech_detail_run = tech_para.add_run(highlights["tech_stack"])
            tech_detail_run.font.size = Pt(9)
            tech_detail_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            tech_para.paragraph_format.space_after = Pt(4)

        # 要点
        if "bullets" in highlights:
            for bullet in highlights["bullets"]:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.size = Pt(10)
                bullet_para.paragraph_format.space_after = Pt(2)

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)
        print(f"简历已保存为: {filename}")


def create_optimized_resume():
    """创建优化后的简历"""
    generator = ResumeGenerator()

    # ============== 联系信息 ==============
    generator.add_name_header(
        name="Your Name",
        phone="1XX-XXXX-XXXX",
        email="your-email@example.com",
        github="https://github.com/your-github"  # 请替换为实际 GitHub
    )

    # ============== 教育背景 ==============
    generator.add_education(
        school="武昌首义学院",
        major="计算机科学与技术（华为云班）",
        degree="本科在读",
        start_date="2023.09",
        end_date="2027.06"
    )

    # ============== 专业技能 ==============
    skills = {
        "AI 后端与并发": [
            "Python asyncio 非阻塞编程",
            "协程并发与线程池隔离调优",
            "HTTP SSE 流式传输",
            "TCP 背压机制"
        ],
        "RAG 与检索优化": [
            "双阶段混合召回 (Dense/Sparse)",
            "RRF 倒数排名融合",
            "Cross-Encoder 精排",
            "向量语义缓存 (VSS)"
        ],
        "Agent 与智能体": [
            "状态机 (FSM) 架构",
            "ReAct 确定性运行时",
            "Tool Calling 机制",
            "任务规划与迭代控制"
        ],
        "文档解析与 ETL": [
            "多模态文档解析",
            "长上下文切分策略",
            "表格防断裂机制",
            "VLM 图像语义提取"
        ],
        "云原生与中间件": [
            "FastAPI",
            "Redis / Qdrant / Neo4j",
            "Celery 分布式调度",
            "Docker / MinIO"
        ]
    }
    generator.add_skills(skills)

    # ============== 项目经历 1: KGateway ==============
    generator.add_experience(
        company="KGateway - 企业级混合多模态 Agent 知识库网关",
        position="AI 后端研发",
        start_date="2026.04",
        end_date="2026.05",
        highlights={
            "description": "面向企业大模型落地场景的 AI 基础设施网关，基于 FastAPI 构建流式异步管道，解决长尾延迟、Token 成本高及多租户隔离问题。",
            "tech_stack": "FastAPI, Qdrant, Redis, Neo4j, BGE-Reranker, LangFuse, Locust",
            "bullets": [
                "设计 HTTP SSE 双向解耦流式网关，实现「状态流 + 文本流」架构；首创 asyncio.wait(FIRST_COMPLETED) 双路竞速守护机制，在模型思考期维持 200ms 高频客户端心跳检测，异常离线时强杀连接，有效阻止 Token 盗刷。",
                "基于 Qdrant HNSW 索引原生 Filter 实现 AND 逻辑预过滤位图剪枝，在索引遍历阶段完成租户数据硬隔离，规避后过滤导致的召回率衰减问题。",
                "构建 Dense/Sparse 双阶段混排管线：采用 2-gram 倒排索引 + RRF 融合算法（k=60）动态排序；通过 asyncio.to_thread 安全卸载 BGE-Reranker 精排推理，避免阻塞主事件循环。",
                "摒弃黑盒框架，基于 FSM 状态机自研确定性 Agent 运行时（Planner → Tool Executor → Fallback），注入最大 4 次迭代沙箱，彻底杜绝无限循环风险。",
                "实现 RediSearch 向量语义缓存，对语义相似度 >0.96 的请求实现 5ms 毫秒级拦截，Token 成本降低 42.6%；手写三态自修复熔断器防御下游 API 速率雪崩。"
            ]
        }
    )

    # ============== 项目经历 2: OmniParse ==============
    generator.add_experience(
        company="OmniParse - 企业级异步多模态文档解析与 ETL 引擎",
        position="AI 数据架构",
        start_date="2026.03",
        end_date="2026.04",
        highlights={
            "description": "面向大模型异构知识库场景的分布式数据清洗与入库流水线，支持复杂 PDF、跨页表格等多源数据，提升脏数据结构化提取的准确率。",
            "tech_stack": "Celery, Redis, MinIO, Unstructured, Qdrant, Docker Compose",
            "bullets": [
                "设计 FastAPI + Celery 异步解耦架构，将耗时文件解析从主业务分离；实现流式写入 MinIO 对象存储并生成临时签名直链，高并发场景下任务提交延迟降至毫秒级。",
                "自研 EnterpriseChunker 格式感知切分器：针对 Token 暴力切分导致的表格语义破碎问题，实现跨页表格整体还原为 Parent Doc，生成摘要 Child Chunk 并保留指针关系，杜绝碎表引发的模型幻觉。",
                "在结构化解析阶段透传 PDF 内嵌图片至轻量级 VLM 生成 Image Caption，完成富文本原位合并，最大程度保留长上下文图文关联性。",
                "设计全流式 Generator Pipeline，将 Qdrant 入库控制在 100 points/batch，高并发场景空间复杂度压制至 O(1)；基于 Docker Compose 实现多级 Worker 一键部署。"
            ]
        }
    )

    # 保存文件
    output_path = os.path.join(os.path.dirname(__file__), "Resume.docx")
    generator.save(output_path)


if __name__ == "__main__":
    create_optimized_resume()
    print("简历生成完成！")
